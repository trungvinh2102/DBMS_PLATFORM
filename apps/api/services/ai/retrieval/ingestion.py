"""
ingestion.py

Document ingestion helpers that extract plain text from files and URLs before
delegating to the existing RAG text indexing pipeline.
"""

import ipaddress
import os
import re
import socket
from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
from typing import Optional
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from .index_service import RagIndexService, rag_index_service


class IngestionError(ValueError):
    """Raised when a source cannot be safely fetched or parsed."""


@dataclass(frozen=True)
class ExtractedDocument:
    """Plain text extracted from an ingestion source."""

    title: str
    content: str
    uri: Optional[str]
    source_type: str


class RagIngestionService:
    """Fetches/parses external sources and indexes their extracted text."""

    def __init__(self, index_service: Optional[RagIndexService] = None):
        self.index_service = index_service or rag_index_service

    def ingest_url(
        self,
        url: str,
        title: Optional[str] = None,
        database_id: Optional[str] = None,
        user_id: Optional[str] = None,
        source_id: Optional[str] = None,
        access_scope: str = "user",
    ) -> dict:
        """Fetches an HTTP(S) URL and indexes the extracted text."""
        normalized_url = self._normalize_url(url)
        payload, content_type, final_url = self._fetch_url(normalized_url)
        extracted = self.extract_bytes(
            payload,
            filename=self._filename_from_url(final_url),
            content_type=content_type,
            title=title or self._title_from_url(final_url),
            uri=final_url,
            source_type="web_page",
        )
        return self.index_service.index_text_source(
            extracted.source_type,
            extracted.title,
            extracted.content,
            database_id=database_id,
            user_id=user_id,
            uri=extracted.uri,
            source_id=source_id,
            access_scope=access_scope,
        )

    def ingest_file(
        self,
        payload: bytes,
        filename: str,
        content_type: Optional[str] = None,
        title: Optional[str] = None,
        database_id: Optional[str] = None,
        user_id: Optional[str] = None,
        source_id: Optional[str] = None,
        access_scope: str = "user",
    ) -> dict:
        """Extracts text from an uploaded file and indexes it as a document."""
        extracted = self.extract_bytes(
            payload,
            filename=filename,
            content_type=content_type,
            title=title or self._title_from_filename(filename),
            uri=f"upload://{filename}",
            source_type="document",
        )
        return self.index_service.index_text_source(
            extracted.source_type,
            extracted.title,
            extracted.content,
            database_id=database_id,
            user_id=user_id,
            uri=extracted.uri,
            source_id=source_id,
            access_scope=access_scope,
        )

    def extract_bytes(
        self,
        payload: bytes,
        filename: str,
        content_type: Optional[str],
        title: str,
        uri: Optional[str],
        source_type: str,
    ) -> ExtractedDocument:
        """Converts supported source bytes into text for downstream chunking."""
        self._validate_payload_size(payload)
        extension = self._extension(filename)
        normalized_type = str(content_type or "").split(";", 1)[0].strip().lower()

        if extension == ".pdf" or normalized_type == "application/pdf":
            content = self._extract_pdf(payload)
        elif extension == ".docx" or normalized_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            content = self._extract_docx(payload)
        elif extension in {".html", ".htm"} or normalized_type == "text/html":
            content = extract_html_text(self._decode_text(payload))
        elif extension in {".md", ".markdown", ".txt", ""} or normalized_type.startswith("text/"):
            content = self._decode_text(payload)
        else:
            raise IngestionError(f"Unsupported ingestion file type: {extension or normalized_type or 'unknown'}")

        cleaned = self._clean_extracted_text(content)
        if not cleaned:
            raise IngestionError("No extractable text found in source")
        return ExtractedDocument(title=title, content=cleaned, uri=uri, source_type=source_type)

    def _fetch_url(self, url: str) -> tuple[bytes, str, str]:
        self._validate_url(url)
        request = Request(url, headers={"User-Agent": "QurioDB-RAG-Ingestion/1.0"})
        try:
            with urlopen(request, timeout=self._url_timeout_seconds()) as response:
                final_url = response.geturl()
                self._validate_url(final_url)
                content_type = response.headers.get("content-type", "")
                payload = response.read(self._max_source_bytes() + 1)
        except (HTTPError, URLError, TimeoutError, OSError) as exc:
            raise IngestionError(f"Failed to fetch URL: {exc}") from exc

        self._validate_payload_size(payload)
        return payload, content_type, final_url

    def _validate_url(self, url: str) -> None:
        parsed = urlparse(url)
        if parsed.scheme not in {"http", "https"} or not parsed.hostname:
            raise IngestionError("Only http and https URLs are supported")
        if self._allow_private_urls():
            return
        try:
            addresses = socket.getaddrinfo(parsed.hostname, parsed.port or (443 if parsed.scheme == "https" else 80))
        except OSError as exc:
            raise IngestionError(f"Unable to resolve URL host: {parsed.hostname}") from exc
        for address in addresses:
            ip = ipaddress.ip_address(address[4][0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast or ip.is_unspecified:
                raise IngestionError("Private, local, or reserved URL hosts are blocked for ingestion")

    def _validate_payload_size(self, payload: bytes) -> None:
        if len(payload or b"") > self._max_source_bytes():
            raise IngestionError(f"Source is larger than {self._max_source_bytes()} bytes")

    def _extract_pdf(self, payload: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise IngestionError("PDF ingestion requires the pypdf package") from exc

        reader = PdfReader(BytesIO(payload))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    def _extract_docx(self, payload: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise IngestionError("DOCX ingestion requires the python-docx package") from exc

        document = Document(BytesIO(payload))
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    blocks.append(" | ".join(values))
        return "\n".join(blocks)

    def _decode_text(self, payload: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "cp1252"):
            try:
                return payload.decode(encoding)
            except UnicodeDecodeError:
                continue
        return payload.decode("utf-8", errors="replace")

    def _normalize_url(self, url: str) -> str:
        github_blob = re.match(r"^https://github\.com/([^/]+)/([^/]+)/blob/([^/]+)/(.+)$", str(url or "").strip())
        if github_blob:
            owner, repo, branch, path = github_blob.groups()
            return f"https://raw.githubusercontent.com/{owner}/{repo}/{branch}/{path}"
        return str(url or "").strip()

    def _clean_extracted_text(self, text: str) -> str:
        lines = [re.sub(r"\s+", " ", line).strip() for line in str(text or "").splitlines()]
        return "\n".join(line for line in lines if line).strip()

    def _filename_from_url(self, url: str) -> str:
        path = urlparse(url).path.rstrip("/")
        return path.rsplit("/", 1)[-1] or "web-page.html"

    def _title_from_url(self, url: str) -> str:
        filename = self._filename_from_url(url)
        return self._title_from_filename(filename) or urlparse(url).hostname or "Web page"

    def _title_from_filename(self, filename: str) -> str:
        name = str(filename or "Document").rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
        return re.sub(r"\.[A-Za-z0-9]+$", "", name).replace("_", " ").replace("-", " ").strip() or "Document"

    def _extension(self, filename: str) -> str:
        name = str(filename or "").lower()
        return f".{name.rsplit('.', 1)[-1]}" if "." in name else ""

    def _max_source_bytes(self) -> int:
        try:
            configured = int(os.getenv("QURIODB_RAG_INGEST_MAX_BYTES", str(10 * 1024 * 1024)))
        except ValueError:
            configured = 10 * 1024 * 1024
        return max(1, configured)

    def _url_timeout_seconds(self) -> int:
        try:
            configured = int(os.getenv("QURIODB_RAG_INGEST_URL_TIMEOUT", "10"))
        except ValueError:
            configured = 10
        return max(1, configured)

    def _allow_private_urls(self) -> bool:
        return os.getenv("QURIODB_RAG_ALLOW_PRIVATE_URLS", "false").lower() in {"1", "true", "yes"}


class _TextHTMLParser(HTMLParser):
    """Small stdlib HTML-to-text parser for local desktop ingestion."""

    SKIPPED_TAGS = {"script", "style", "noscript", "svg"}
    BLOCK_TAGS = {"address", "article", "aside", "blockquote", "br", "div", "footer", "h1", "h2", "h3", "h4", "h5", "h6", "header", "li", "main", "p", "section", "table", "tr"}

    def __init__(self):
        super().__init__()
        self.parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self.SKIPPED_TAGS:
            self._skip_depth += 1
        if tag in self.BLOCK_TAGS:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self.SKIPPED_TAGS and self._skip_depth:
            self._skip_depth -= 1
        if tag in self.BLOCK_TAGS:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip_depth and data.strip():
            self.parts.append(data.strip())

    def text(self) -> str:
        return "\n".join(part.strip() for part in self.parts if part.strip())


def extract_html_text(html: str) -> str:
    """Extracts visible text from HTML without external parser dependencies."""
    parser = _TextHTMLParser()
    parser.feed(str(html or ""))
    return parser.text()


rag_ingestion_service = RagIngestionService()
